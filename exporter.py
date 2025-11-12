"""
Export translated chapters to various formats (DOCX, EPUB, PDF).
"""
from pathlib import Path
from typing import List
from loguru import logger

from models import Chapter
from config import settings


def export_to_docx(chapters: List[Chapter], output_file: Path):
    """Export chapters to Microsoft Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return
    
    logger.info(f"Exporting to Word document: {output_file}")
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Translated Novel', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata
    doc.add_paragraph(f'Total Chapters: {len(chapters)}')
    doc.add_paragraph('')
    
    # Add each chapter
    for chapter in chapters:
        # Chapter title
        doc.add_heading(chapter.title_vietnamese, level=1)
        
        # Original title
        original = doc.add_paragraph(f'Original: {chapter.title_chinese}')
        original.italic = True
        
        doc.add_paragraph('')
        
        # Chapter content - split by paragraphs
        paragraphs = chapter.content_vietnamese.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                # Set font
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
        
        # Add page break after each chapter
        doc.add_page_break()
    
    # Save document
    doc.save(output_file)
    logger.success(f"Word document exported to {output_file}")


def export_to_epub(chapters: List[Chapter], output_file: Path):
    """Export chapters to EPUB format."""
    try:
        from ebooklib import epub
    except ImportError:
        logger.error("ebooklib not installed. Run: pip install ebooklib")
        return
    
    logger.info(f"Exporting to EPUB: {output_file}")
    
    # Create book
    book = epub.EpubBook()
    
    # Set metadata
    book.set_identifier('translated_novel_001')
    book.set_title('Translated Novel')
    book.set_language('vi')
    book.add_author('Unknown Author')
    
    # Create chapters
    epub_chapters = []
    spine = ['nav']
    
    for i, chapter in enumerate(chapters, 1):
        # Create chapter
        c = epub.EpubHtml(
            title=chapter.title_vietnamese,
            file_name=f'chapter_{i:04d}.xhtml',
            lang='vi'
        )
        
        # Format content
        content = f'''
        <h1>{chapter.title_vietnamese}</h1>
        <p><em>Original: {chapter.title_chinese}</em></p>
        <br/>
        '''
        
        # Add paragraphs
        paragraphs = chapter.content_vietnamese.split('\n\n')
        for para in paragraphs:
            if para.strip():
                content += f'<p>{para.strip()}</p>\n'
        
        c.content = content
        
        # Add chapter to book
        book.add_item(c)
        epub_chapters.append(c)
        spine.append(c)
    
    # Define Table of Contents
    book.toc = tuple(epub_chapters)
    
    # Add navigation files
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    # Define CSS style
    style = '''
    @namespace epub "http://www.idpf.org/2007/ops";
    body {
        font-family: Times, Times New Roman, serif;
        font-size: 1.1em;
        line-height: 1.6;
        margin: 5%;
    }
    h1 {
        text-align: center;
        margin-top: 2em;
        margin-bottom: 1em;
    }
    p {
        text-align: justify;
        text-indent: 1.5em;
        margin: 0.5em 0;
    }
    em {
        color: #666;
    }
    '''
    
    nav_css = epub.EpubItem(
        uid="style_nav",
        file_name="style/nav.css",
        media_type="text/css",
        content=style
    )
    book.add_item(nav_css)
    
    # Set spine
    book.spine = spine
    
    # Write to file
    epub.write_epub(output_file, book)
    logger.success(f"EPUB exported to {output_file}")


def export_to_pdf(chapters: List[Chapter], output_file: Path):
    """Export chapters to PDF format."""
    try:
        from weasyprint import HTML, CSS
        from markdown import markdown
    except ImportError:
        logger.error("weasyprint or markdown not installed. Run: pip install weasyprint markdown")
        return
    
    logger.info(f"Exporting to PDF: {output_file}")
    
    # Generate HTML content
    html_content = '''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: "Times New Roman", Times, serif;
                font-size: 12pt;
                line-height: 1.6;
                text-align: justify;
            }
            h1 {
                text-align: center;
                font-size: 24pt;
                margin-top: 2em;
                margin-bottom: 0.5em;
                page-break-before: always;
            }
            h1:first-of-type {
                page-break-before: auto;
            }
            .original-title {
                text-align: center;
                font-style: italic;
                color: #666;
                margin-bottom: 2em;
            }
            p {
                text-indent: 1.5em;
                margin: 0.5em 0;
            }
            .book-title {
                text-align: center;
                font-size: 32pt;
                margin-top: 3em;
                margin-bottom: 0.5em;
                page-break-after: always;
            }
            .metadata {
                text-align: center;
                color: #666;
                margin-bottom: 2em;
            }
        </style>
    </head>
    <body>
        <div class="book-title">Translated Novel</div>
        <div class="metadata">Total Chapters: ''' + str(len(chapters)) + '''</div>
    '''
    
    # Add each chapter
    for chapter in chapters:
        html_content += f'''
        <h1>{chapter.title_vietnamese}</h1>
        <p class="original-title">Original: {chapter.title_chinese}</p>
        '''
        
        # Add paragraphs
        paragraphs = chapter.content_vietnamese.split('\n\n')
        for para in paragraphs:
            if para.strip():
                html_content += f'<p>{para.strip()}</p>\n'
    
    html_content += '''
    </body>
    </html>
    '''
    
    # Convert HTML to PDF
    HTML(string=html_content).write_pdf(output_file)
    logger.success(f"PDF exported to {output_file}")


def export_to_markdown(chapters: List[Chapter], output_file: Path):
    """Export chapters as Markdown file."""
    with output_file.open('w', encoding='utf-8') as f:
        # Write header
        f.write("# Translated Novel\n\n")
        f.write(f"Total Chapters: {len(chapters)}\n\n")
        f.write("---\n\n")
        
        # Write each chapter
        for chapter in chapters:
            f.write(f"## {chapter.title_vietnamese}\n\n")
            f.write(f"*Original: {chapter.title_chinese}*\n\n")
            f.write(chapter.content_vietnamese)
            f.write("\n\n---\n\n")
    
    logger.success(f"Markdown exported to {output_file}")
